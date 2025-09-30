#!/usr/bin/env python3
"""
Script to create a DOCX document explaining the Note Ninjas query flow
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
import os

def add_heading_with_emoji(doc, text, level=1):
    """Add heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_code_block(doc, code, language="python"):
    """Add formatted code block"""
    para = doc.add_paragraph()
    para.style = doc.styles['Normal']
    
    # Set monospace font
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # Add the code text
    para.add_run(code)
    
    # Set paragraph formatting
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_after = Pt(6)
    
    return para

def add_json_block(doc, json_text):
    """Add formatted JSON block"""
    para = doc.add_paragraph()
    
    # Add JSON label
    label_run = para.add_run("JSON Response:\n")
    label_run.bold = True
    label_run.font.size = Pt(11)
    
    # Add JSON content
    json_run = para.add_run(json_text)
    json_run.font.name = 'Courier New'
    json_run.font.size = Pt(9)
    
    # Set paragraph formatting
    para_format = para.paragraph_format
    para_format.left_indent = Inches(0.5)
    para_format.space_after = Pt(12)
    
    return para

def add_flowchart_step(doc, step_number, title, description):
    """Add a flowchart-style step"""
    para = doc.add_paragraph()
    
    # Step number and title
    step_run = para.add_run(f"STEP {step_number}: {title}")
    step_run.bold = True
    step_run.font.size = Pt(14)
    step_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    
    para.add_run("\n")
    
    # Description
    desc_run = para.add_run(description)
    desc_run.font.size = Pt(11)
    
    para_format = para.paragraph_format
    para_format.space_after = Pt(12)
    
    return para

def create_query_flow_document():
    """Create the main document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('🍼 Note Ninjas Backend: Complete Query Flow Explanation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('A Step-by-Step Journey from Input to Output')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_page_break()
    
    # Table of Contents
    add_heading_with_emoji(doc, "📋 Table of Contents", 1)
    toc_items = [
        "1. Overview & System Architecture",
        "2. Complete Query Flow (10 Steps)",
        "3. Example Request & Response",
        "4. Code References",
        "5. Visual Flow Summary"
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item, style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # Overview
    add_heading_with_emoji(doc, "🏗️ System Overview", 1)
    
    overview_text = """
The Note Ninjas Backend is a RAG-only (Retrieval-Augmented Generation) recommendation engine specifically designed for occupational therapy professionals. Think of it as a super-smart medical librarian that:

• Never makes up information (no hallucination)
• Only uses verified medical documents
• Provides structured, evidence-based recommendations
• Learns from user feedback over time
• Includes proper medical citations for everything

The system processes medical documents at startup and creates a searchable knowledge base. When you send a query, it finds the most relevant information and uses GPT-4 to organize it into helpful recommendations.
    """.strip()
    
    doc.add_paragraph(overview_text)
    
    # Architecture diagram (text-based)
    add_heading_with_emoji(doc, "🎯 Architecture Flow", 2)
    
    arch_text = """
    Your Query
        ↓
    [Feedback Check] → Remember user preferences
        ↓
    [Query Building] → Convert input to search terms
        ↓
    [Hybrid Retrieval] → BM25 + Vector Search
        ↓
    [Reranking] → Cross-encoder scoring
        ↓
    [Context Building] → Prepare research packet
        ↓
    [GPT-4 Generation] → Create structured response
        ↓
    [Parsing] → Convert to final format
        ↓
    Structured Recommendations
    """
    
    para = doc.add_paragraph(arch_text)
    para.runs[0].font.name = 'Courier New'
    para.runs[0].font.size = Pt(11)
    para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # Detailed Steps
    add_heading_with_emoji(doc, "🔍 Complete Query Flow - Step by Step", 1)
    
    doc.add_paragraph("Let's trace exactly what happens when you send this example query:")
    
    # Example request
    example_request = '''{
  "user_input": {
    "patient_condition": "21 year old female with torn rotator cuff",
    "desired_outcome": "increase right shoulder abduction to 150°",
    "treatment_progression": "progressed from 130° to 135°"
  },
  "session_id": "session_123"
}'''
    
    add_json_block(doc, example_request)
    
    # Step 1
    add_flowchart_step(doc, 1, "🎬 Query Arrives", 
        "Your HTTP POST request hits the /recommendations endpoint in main.py. FastAPI validates your JSON data using Pydantic models and logs the session ID.")
    
    add_code_block(doc, '''# main.py line 84-106
@app.post("/recommendations", response_model=RecommendationResponse)
async def get_recommendations(
    request: RecommendationRequest,
    rag: GPTRAGSystem = Depends(get_rag_system),
    feedback: FeedbackManager = Depends(get_feedback_manager)
):
    logger.info(f"Processing request for session: {request.session_id}")
    
    feedback_state = feedback.get_feedback_state(request.session_id)
    
    recommendations = await rag.generate_recommendations(
        user_input=request.user_input,
        rag_manifest=request.rag_manifest,
        session_id=request.session_id,
        feedback_state=feedback_state
    )
    
    return recommendations''')
    
    # Step 2
    add_flowchart_step(doc, 2, "🧠 Check User Memory", 
        "The system checks if you've used it before. It looks for preferences like 'make exercises easier' or 'don't use CPT code 97530'. If it's your first time, it creates an empty feedback state.")
    
    doc.add_paragraph("Think of it like: The librarian remembers you don't like really hard books! 📚")
    
    # Step 3
    add_flowchart_step(doc, 3, "🚀 RAG Process Starts", 
        "The main generate_recommendations function in gpt_rag_system.py begins. This is where all the AI magic happens!")
    
    add_code_block(doc, '''# gpt_rag_system.py line 148-173
async def generate_recommendations(
    self,
    user_input: UserInput,
    rag_manifest: RAGManifest,
    session_id: str,
    feedback_state: Optional[Dict[str, Any]] = None
) -> RecommendationResponse:
    # Build search query from user input
    query = self._build_query(user_input)
    
    # Search the medical knowledge base
    retrieval_results = self.retriever.search(
        query=query,
        top_k=rag_manifest.max_sources
    )
    
    # Improve results with AI reranking
    reranked_results = self.reranker.rerank(
        query=query,
        results=retrieval_results,
        top_n=min(rag_manifest.max_sources, 12)
    )
    
    # Prepare context and generate response
    context = self._prepare_context(reranked_results, user_input, feedback_state)
    response = await self._generate_gpt_response(context, user_input, rag_manifest)
    return self._parse_gpt_response(response)''')
    
    # Step 4
    add_flowchart_step(doc, 4, "📝 Build Search Query", 
        "Your input gets converted into a search query. The system combines your patient condition, desired outcome, and treatment progression into one search string.")
    
    query_example = '''Input: "21 year old female with torn rotator cuff" + "increase right shoulder abduction to 150°"
Output: "21 year old female with torn rotator cuff increase right shoulder abduction to 150°"'''
    
    para = doc.add_paragraph(query_example)
    para.runs[0].font.name = 'Courier New'
    para.runs[0].font.size = Pt(10)
    para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph("Think of it like: The librarian writes down exactly what you're looking for! 🔍")
    
    # Step 5
    add_flowchart_step(doc, 5, "🔍 Search the Medical Library", 
        "This is where the hybrid retrieval happens. The system uses TWO different search methods simultaneously:")
    
    doc.add_paragraph("5a) BM25 Search (Keyword matching):")
    bm25_text = "• Looks for exact words like 'rotator cuff' and 'shoulder abduction'\n• Finds documents that mention these exact terms\n• Fast but only finds exact matches"
    doc.add_paragraph(bm25_text)
    
    doc.add_paragraph("5b) Vector Search (Meaning matching):")
    vector_text = "• Uses OpenAI embeddings to understand MEANING\n• Finds documents about similar concepts even with different words\n• Like finding 'shoulder rehabilitation' when you said 'shoulder therapy'"
    doc.add_paragraph(vector_text)
    
    doc.add_paragraph("5c) Combine Results:")
    combine_text = "• Merges both search results\n• Prioritizes Note Ninjas documents over textbooks\n• Returns top 50 most relevant chunks"
    doc.add_paragraph(combine_text)
    
    doc.add_paragraph("Think of it like: The librarian searches both the card catalog AND asks other librarians! 📚🔍")
    
    # Step 6
    add_flowchart_step(doc, 6, "🏆 Rank Results (Reranking)", 
        "A cross-encoder AI model scores how well each document matches your specific query. It's like having an expert read each document and rate its usefulness from 1-10.")
    
    rerank_process = '''Process:
1. Take your query + each document
2. Cross-encoder scores relevance (0-1)
3. Sort from most relevant to least relevant
4. Remove duplicate information
5. Return top 12 most relevant documents'''
    
    para = doc.add_paragraph(rerank_process)
    para.runs[0].font.name = 'Courier New'
    para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph("Think of it like: The librarian ranks books from 'most helpful' to 'least helpful'! 📊")
    
    # Step 7
    add_flowchart_step(doc, 7, "📋 Prepare Context for AI", 
        "The system creates a comprehensive research packet containing your question, your preferences, and all the relevant medical document excerpts.")
    
    context_example = '''USER INPUT:
Patient Condition: 21 year old female with torn rotator cuff
Desired Outcome: increase right shoulder abduction to 150°

RETRIEVED SOURCES:
Source 1:
Type: note_ninjas
File: Shoulder_Rehabilitation
Headers: Rotator Cuff Exercises
Page: p. 23
Content: For rotator cuff tears, begin with passive range of motion exercises...

Source 2:
Type: cpg
File: APTA_Shoulder_Guidelines
Content: Evidence suggests progressive loading for rotator cuff repairs...'''
    
    add_code_block(doc, context_example)
    
    doc.add_paragraph("Think of it like: The librarian makes you a custom research packet! 📑")
    
    # Step 8
    add_flowchart_step(doc, 8, "🤖 Ask GPT-4 for Recommendations", 
        "The system sends everything to GPT-4o Mini with very strict instructions. The AI can ONLY use information from the provided sources - it cannot make anything up!")
    
    gpt_settings = '''GPT-4 Configuration:
• Model: gpt-4o-mini
• Temperature: 0.1 (very focused, not creative)
• Max tokens: 2000
• Response format: JSON only
• System prompt: Detailed medical AI instructions (94 lines!)

Key Rules for GPT-4:
• Only use information from provided sources
• Include citations for everything
• Don't fabricate CPT codes
• Use clinical language
• Return structured JSON'''
    
    para = doc.add_paragraph(gpt_settings)
    para.runs[0].font.name = 'Courier New'
    para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph("Think of it like: The librarian writes you a perfect medical report following strict rules! 🤖📋")
    
    # Step 9
    add_flowchart_step(doc, 9, "📊 Parse GPT Response", 
        "GPT-4 returns pure JSON. The system converts this into proper Python objects with validation and error checking.")
    
    json_structure = '''{
  "high_level": [
    "Begin with pain-free passive ROM, progress weekly by 5-10°"
  ],
  "subsections": [
    {
      "title": "Range of Motion Restoration",
      "rationale": "Progressive loading supports tissue healing",
      "exercises": [
        {
          "title": "Passive Shoulder Abduction",
          "description": "Patient supine, therapist moves arm through pain-free range",
          "cues": ["Relax the shoulder", "Let me do the work"],
          "documentation": "Instructed passive shoulder abduction 0-135°...",
          "cpt": "97140",
          "sources": [
            {
              "type": "note_ninjas",
              "id": "Shoulder_Rehabilitation",
              "section": "Passive ROM",
              "page": "p. 23",
              "quote": "Passive abduction should be performed..."
            }
          ]
        }
      ]
    }
  ],
  "confidence": "high"
}'''
    
    add_json_block(doc, json_structure)
    
    doc.add_paragraph("Think of it like: The librarian organizes the report into neat sections! 🗂️")
    
    # Step 10
    add_flowchart_step(doc, 10, "✅ Return Final Response", 
        "The structured recommendations are sent back to you as JSON. You get detailed exercises, CPT codes, safety considerations, and citations for everything!")
    
    final_features = '''What you receive:
✓ High-level treatment recommendations
✓ Detailed exercise instructions with cues
✓ CPT billing codes (only verified ones)
✓ Clinical documentation examples
✓ Safety considerations and contraindications
✓ Source citations for every recommendation
✓ Confidence level (high/medium/low)
✓ Alternative approaches when appropriate'''
    
    para = doc.add_paragraph(final_features)
    para.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # Feedback System
    add_heading_with_emoji(doc, "🔄 Bonus: Feedback Learning", 1)
    
    feedback_text = '''When you provide feedback like "too advanced" or correct a CPT code:

1. The system stores your preference in the FeedbackManager
2. Next time you make a request, it remembers your preferences
3. Adjusts future recommendations based on your feedback
4. Builds a personalized profile over your session

Example feedback types:
• Thumbs up/down → Influences future ranking
• "Too advanced" → System suggests easier exercises
• CPT corrections → Learns correct billing codes
• Content blocking → Removes inappropriate recommendations

This creates a personalized experience that improves over time!'''
    
    doc.add_paragraph(feedback_text)
    
    # Magic Summary
    add_heading_with_emoji(doc, "💡 The Magic Summary", 1)
    
    summary_flow = '''
Your Question → Check Memory → Build Search → Find Documents → Rank by Relevance 
→ Create Research Packet → Ask GPT-4 → Parse Response → Return Structured Recommendations
    '''
    
    para = doc.add_paragraph(summary_flow.strip())
    para.runs[0].font.name = 'Courier New'
    para.runs[0].font.size = Pt(12)
    para.runs[0].bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    key_insight = '''🎯 KEY INSIGHT: Your system NEVER makes up medical information! 

Everything comes from real medical documents that were processed during startup. GPT-4 is just a very smart organizer that puts the right information together in a helpful format.

Think of it like: A super-powered medical librarian who instantly finds perfect research, reads it all, and writes you a custom treatment plan! 🏥📚🤖'''
    
    doc.add_paragraph(key_insight)
    
    # Code References
    doc.add_page_break()
    add_heading_with_emoji(doc, "📂 Key Code File References", 1)
    
    code_refs = '''
main.py → API endpoints and request handling
├── Lines 84-106: Main /recommendations endpoint
├── Lines 93: Feedback state retrieval
├── Lines 95-100: RAG system call

core/gpt_rag_system.py → Main RAG orchestration
├── Lines 148-173: generate_recommendations() main flow
├── Lines 175-185: Query building from user input
├── Lines 187-217: Context preparation for GPT-4
├── Lines 253-270: GPT-4 API call with system prompt
├── Lines 272-334: JSON response parsing

core/retriever.py → Hybrid search system
├── BM25 keyword search implementation
├── OpenAI embeddings integration
├── Result combination and source boosting

core/reranker.py → Cross-encoder reranking
├── Cross-encoder model loading
├── Query-document pair scoring
├── Diversity filtering

core/feedback_manager.py → Learning system
├── Feedback storage and processing
├── Preference extraction from user input
├── Session state management

models/ → Data structures
├── request_models.py: Input validation
├── response_models.py: Output formatting
    '''
    
    para = doc.add_paragraph(code_refs.strip())
    para.runs[0].font.name = 'Courier New'
    para.runs[0].font.size = Pt(10)
    
    # Final page
    doc.add_page_break()
    add_heading_with_emoji(doc, "🎓 Conclusion", 1)
    
    conclusion = '''The Note Ninjas Backend represents a sophisticated application of modern AI techniques specifically tailored for medical professionals. By combining retrieval-augmented generation with domain-specific knowledge, it provides:

• Evidence-based recommendations without hallucination
• Proper medical citations and billing codes
• Personalized learning through feedback
• Professional-grade safety and accuracy

Every query follows this exact 10-step process, ensuring consistent, reliable, and medically sound recommendations for occupational therapy professionals.

The system bridges the gap between vast medical literature and practical clinical application, serving as an intelligent assistant that enhances professional decision-making while maintaining the highest standards of medical accuracy and safety.'''
    
    doc.add_paragraph(conclusion)
    
    # Add footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Generated by Note Ninjas Documentation System")
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

def main():
    """Main function to create and save the document"""
    print("🔄 Creating Note Ninjas Query Flow Documentation...")
    
    try:
        doc = create_query_flow_document()
        
        # Save the document
        output_path = "Note_Ninjas_Query_Flow_Explained.docx"
        doc.save(output_path)
        
        print(f"✅ Document created successfully!")
        print(f"📄 Saved as: {output_path}")
        print(f"📍 Location: {os.path.abspath(output_path)}")
        
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())